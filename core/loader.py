"""Extract plain text from documents (PDF, DOCX, TXT, MD, CSV, JPG, JPEG, PNG)."""
from __future__ import annotations

from pathlib import Path

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".csv", ".jpg", ".jpeg", ".png")

# ------------------------------------------------------------------ OCR hook
# Registered by the service layer once the Gemini client is available.
_ocr_client = None
_ocr_model: str = "gemini-2.0-flash"


def register_ocr_client(client, model: str | None = None) -> None:
    global _ocr_client, _ocr_model
    _ocr_client = client
    if model:
        _ocr_model = model


def _gemini_ocr(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """Send raw image bytes to Gemini vision and return extracted text."""
    from google.genai import types
    response = _ocr_client.models.generate_content(
        model=_ocr_model,
        contents=types.Content(parts=[
            types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
            types.Part.from_text(
                text=(
                    "Extract all text from this document image exactly as it appears. "
                    "Preserve numbers, labels, and layout. Return only the extracted text."
                )
            ),
        ]),
    )
    return (response.text or "").strip()


# ----------------------------------------------------------------- public API

def load_document(path: str | Path) -> list[dict]:
    """Return a list of text blocks: {"page": int | None, "text": str}."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix in (".txt", ".md", ".csv"):
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        return [{"page": None, "text": text}] if text else []
    if suffix in (".jpg", ".jpeg", ".png"):
        return _load_image(path)
    raise ValueError(
        f"Unsupported file type: {suffix} (supported: {', '.join(SUPPORTED_EXTENSIONS)})"
    )


# ----------------------------------------------------------------- loaders

def _load_pdf(path: Path) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    blocks: list[dict] = []
    scanned: list[int] = []  # zero-based page indices with no selectable text

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            blocks.append({"page": i + 1, "text": text})
        else:
            scanned.append(i)

    # Run Gemini OCR on any pages that had no selectable text
    if scanned:
        if _ocr_client is None:
            if not blocks:
                raise ValueError(
                    f"No selectable text in {path.name} — scanned PDF detected. "
                    "OCR is initialising; please try again in a moment."
                )
            # Some pages had text — skip the scanned ones silently
        else:
            import fitz  # pymupdf

            doc = fitz.open(str(path))
            for i in scanned:
                pix = doc[i].get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                text = _gemini_ocr(pix.tobytes("png"))
                if text:
                    blocks.append({"page": i + 1, "text": text})
            doc.close()
            blocks.sort(key=lambda b: b["page"])

    if not blocks:
        raise ValueError(f"Could not extract any text from {path.name}.")
    return blocks


def _load_image(path: Path) -> list[dict]:
    if _ocr_client is None:
        raise ValueError(
            f"Cannot read image {path.name}: OCR client not yet initialised."
        )
    mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png"}
    mime = mime_map.get(path.suffix.lower(), "image/jpeg")
    text = _gemini_ocr(path.read_bytes(), mime)
    if not text:
        raise ValueError(f"Could not extract any text from {path.name}.")
    return [{"page": None, "text": text}]


def _load_docx(path: Path) -> list[dict]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return [{"page": None, "text": text}] if text else []
